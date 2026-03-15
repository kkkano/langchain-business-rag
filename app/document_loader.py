from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


@dataclass
class LoadedDocument:
    source_name: str
    source_path: str
    source_type: str
    text: str


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"[ \t]*\n[ \t]*", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def infer_source_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return "text"
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    raise RuntimeError(f"暂不支持的文档类型: {path.suffix}")


def load_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"无法识别文本编码: {path}")


def load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            pages.append(f"第{page_number}页\n{page_text}")
    return "\n\n".join(pages)


def load_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def load_document(path: Path) -> LoadedDocument:
    actual_path = path.expanduser().resolve()
    if not actual_path.exists():
        raise RuntimeError(f"文件不存在: {actual_path}")
    if actual_path.is_dir():
        raise RuntimeError(f"当前只支持文件，不支持目录: {actual_path}")

    source_type = infer_source_type(actual_path)
    if source_type == "text":
        raw_text = load_text_file(actual_path)
    elif source_type == "pdf":
        raw_text = load_pdf(actual_path)
    elif source_type == "docx":
        raw_text = load_docx(actual_path)
    else:
        raise RuntimeError(f"无法处理的文档类型: {source_type}")

    normalized = normalize_text(raw_text)
    if not normalized:
        raise RuntimeError(f"文档为空，无法建立知识库: {actual_path}")

    return LoadedDocument(
        source_name=actual_path.name,
        source_path=str(actual_path),
        source_type=source_type,
        text=normalized,
    )


def save_uploaded_file(upload_dir: Path, filename: str, content: bytes) -> Path:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise RuntimeError(
            f"不支持的上传类型: {suffix}。当前支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    safe_name = Path(filename).name
    target = upload_dir / safe_name
    counter = 1
    while target.exists():
        target = upload_dir / f"{Path(safe_name).stem}_{counter}{suffix}"
        counter += 1
    target.write_bytes(content)
    return target
